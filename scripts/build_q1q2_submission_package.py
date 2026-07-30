#!/usr/bin/env python3
"""Build guarded Radiology: AI submission previews or the final package.

Preview mode is intentionally watermarked and replaces every claim token with
an explicit pending marker. Final mode accepts only Gate J-rendered sources and
requires passing Gate H/I/J artifacts plus author-confirmed metadata.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "submission"
RADAI = SUBMISSION / "radiology_ai"
MANUSCRIPT = ROOT / "manuscript"
CLAIM_TOKEN = re.compile(
    r"\{\{claim:([A-Z0-9_.-]+)\|"
    r"(?:raw|integer|2f|3f|4f|percent1|percent2|pvalue)\}\}"
)
UNRESOLVED = re.compile(r"\{\{[^{}]+\}\}")
PENDING_MARKER = re.compile(r"\[(?:PENDING|AUTHOR CONFIRMATION REQUIRED|"
                            r"GENERATED DURING FINAL BUILD)[^\]]*\]")
WORD = re.compile(r"\b[\w-]+\b")

PREVIEW_WARNING = (
    "PRE-RESULTS PREVIEW - NOT FOR SUBMISSION. Scientific result fields and "
    "author declarations remain unresolved. Do not upload this file to a "
    "journal."
)

TITLE = (
    "Capacity- and Compute-Matched Evaluation of Published BU-Net Components "
    "with Multi-Seed Development and Independent External Testing for "
    "Multimodal Glioma Segmentation"
)

ABSTRACT_TEMPLATE = """## Abstract

### Purpose

To test whether the published Residual Extended Skip (RES) component retains
a measurable benefit after parameter matching and to characterize uncertainty,
lesion-level failure, external domain shift, and measured resource cost.

### Materials and Methods

This retrospective public-data study used BraTS 2020 for development and
BraTS-Africa for independent external testing. Twelve models were assigned the
same five seeds in five patient-level folds. After development-only model,
loss, checkpoint, endpoint, and analysis freeze, the confirmatory external
cohort was evaluated once. The primary endpoint was patient-level mean whole
tumor, tumor core, and enhancing tumor Dice. The primary contrast compared
U-Net+RES with a parameter-matched plain U-Net using paired bootstrap
intervals, paired sign-flip tests with Holm correction, and hierarchical
seed-patient resampling.

### Results

<!-- BEGIN_ARTIFACT_BOUND_RESULTS -->
The external confirmatory analysis included
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.PAIRED_PATIENT_COUNT|integer}}
paired patients. Mean regional Dice was
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.FIRST_MEAN|3f}}
for U-Net+RES and
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.SECOND_MEAN|3f}}
for the parameter-matched plain U-Net. The paired mean difference was
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.MEAN_DIFFERENCE|3f}}
(95% interval,
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.PAIRED_BOOTSTRAP_LOWER_95|3f}}
to
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.PAIRED_BOOTSTRAP_UPPER_95|3f}});
the Holm-adjusted P value was
{{claim:CONTRAST.UNET_RES_VS_UNET_PARAMETER_MATCHED_RES.HOLM_ADJUSTED_P|pvalue}}.
<!-- END_ARTIFACT_BOUND_RESULTS -->

### Conclusion

The result provides capacity-controlled external evidence about a published
BU-Net component and is bounded to the frozen cohorts, models, endpoints, and
practical interpretation threshold.
"""


def digest(path: Path) -> str:
    value = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            value.update(chunk)
    return value.hexdigest()


def run(command: list[str]) -> None:
    subprocess.run(command, cwd=ROOT, check=True)


def load_json(path: Path) -> dict[str, Any]:
    loaded = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(loaded, dict):
        raise ValueError(f"Expected JSON object: {path}")
    return loaded


def replace_claims_with_pending(source: str) -> str:
    return CLAIM_TOKEN.sub(lambda match: f"[PENDING:{match.group(1)}]", source)


def add_preview_warning(source: str) -> str:
    return f"> **{PREVIEW_WARNING}**\n\n{source}"


def section(source: str, start: str, end: str | None = None) -> str:
    start_match = re.search(rf"^## {re.escape(start)}\s*$", source, re.MULTILINE)
    if start_match is None:
        raise RuntimeError(f"Required section not found: {start}")
    if end is None:
        return source[start_match.start() :]
    end_match = re.search(rf"^## {re.escape(end)}\s*$", source, re.MULTILINE)
    if end_match is None or end_match.start() <= start_match.start():
        raise RuntimeError(f"Required end section not found: {end}")
    return source[start_match.start() : end_match.start()]


def radiology_ai_manuscript(source: str, *, preview: bool) -> str:
    body = section(source, "Introduction")
    body = re.sub(
        r"^## Conclusion\s*$",
        "### Conclusion",
        body,
        flags=re.MULTILINE,
    )
    body = re.sub(
        r"^## (Ethics statement|Funding|Conflicts of interest|"
        r"Author contributions|Use of generative AI and assisted technologies)"
        r"\s*$.*?(?=^## References\s*$)",
        "",
        body,
        flags=re.MULTILINE | re.DOTALL,
    )
    front = f"""# {TITLE}

**Article type:** Original Research

**Summary statement:** [PENDING:SUMMARY_STATEMENT_FROM_FINAL_RESULT]

**Key Points**

1. [PENDING:KEY_POINT_PRIMARY_CAPACITY_CONTROLLED_EFFECT]
2. [PENDING:KEY_POINT_EXTERNAL_FAILURE_OR_DOMAIN_SHIFT]
3. [PENDING:KEY_POINT_RESOURCE_TRADEOFF]

"""
    manuscript = f"{front}{ABSTRACT_TEMPLATE}\n\n{body}"
    manuscript = manuscript.replace(
        "## References",
        "## Anonymized Acknowledgments\n\n"
        "[AUTHOR CONFIRMATION REQUIRED: enter an anonymized acknowledgment or "
        "state none.]\n\n## Figure Legends\n\n"
        "[GENERATED DURING FINAL BUILD: no more than six figures.]\n\n"
        "## References",
        1,
    )
    manuscript += (
        "\n\n## Tables\n\n"
        "[GENERATED DURING FINAL BUILD: no more than four tables; each table "
        "will begin on a separate page and contain no merged cells.]\n"
    )
    if preview:
        manuscript = replace_claims_with_pending(manuscript)
        manuscript = add_preview_warning(manuscript)
    return manuscript


def confirmed_author_metadata(path: Path) -> dict[str, Any]:
    metadata = load_json(path)
    required_scalars = [
        "funding_statement",
        "funder_role",
        "conflict_of_interest_statement",
        "industry_support_statement",
        "ethics_statement",
        "acknowledgments",
        "subject_overlap_statement",
        "generative_ai_disclosure",
        "data_sharing_statement",
    ]
    missing = [
        key for key in required_scalars if not str(metadata.get(key) or "").strip()
    ]
    if metadata.get("confirmed_by_all_authors") is not True:
        missing.append("confirmed_by_all_authors")
    if metadata.get("sole_submission_confirmed") is not True:
        missing.append("sole_submission_confirmed")
    if not metadata.get("authors"):
        missing.append("authors")
    if not metadata.get("affiliations"):
        missing.append("affiliations")
    corresponding = metadata.get("corresponding_author")
    if not isinstance(corresponding, dict) or any(
        not str(corresponding.get(key) or "").strip()
        for key in (
            "name",
            "degrees",
            "email",
            "telephone",
            "street_address",
            "city",
            "postal_code",
            "country",
        )
    ):
        missing.append("corresponding_author")
    if missing:
        raise PermissionError(
            "Author metadata is not complete or confirmed: "
            + ", ".join(sorted(set(missing)))
        )
    return metadata


def gate_is_pass(path: Path) -> bool:
    if not path.is_file():
        return False
    payload = load_json(path)
    return payload.get("status") in {"pass", "complete"}


def require_final_gates() -> None:
    required = [
        ROOT / "artifacts/q1q2_v2/gate_h_completion.json",
        ROOT / "artifacts/q1q2_v2/reproducibility/completion.json",
        ROOT / "artifacts/q1q2_v2/claims/completion.json",
    ]
    failed = [path.as_posix() for path in required if not gate_is_pass(path)]
    if failed:
        raise PermissionError(
            "Final submission build blocked by incomplete gates: "
            + ", ".join(failed)
        )


def author_name(author: dict[str, Any]) -> str:
    parts = [
        str(author.get("given_names") or "").strip(),
        str(author.get("middle_initials") or "").strip(),
        str(author.get("last_name") or "").strip(),
    ]
    name = " ".join(value for value in parts if value)
    degrees = str(author.get("degrees") or "").strip()
    return f"{name}, {degrees}" if degrees else name


def final_title_page(metadata: dict[str, Any], word_count: int) -> str:
    affiliations = {
        str(entry["id"]): entry for entry in metadata["affiliations"]
    }
    author_lines = []
    for author in metadata["authors"]:
        ids = [str(value) for value in author.get("affiliation_ids", [])]
        author_lines.append(
            f"- {author_name(author)}; affiliations {', '.join(ids)}; "
            f"ORCID {author.get('orcid') or 'not provided'}"
        )
    affiliation_lines = [
        f"- {key}: {entry['institution']}, {entry['department']}, "
        f"{entry['city']}, {entry['country']}"
        for key, entry in affiliations.items()
    ]
    corresponding = metadata["corresponding_author"]
    institution = metadata["originating_institution"]
    contributions = "\n".join(
        f"- {entry['author']}: {', '.join(entry['roles'])}"
        for entry in metadata["credit_contributions"]
    )
    return f"""# Full Title Page

## Manuscript

**Title:** {metadata['manuscript_title']}

**Short title:** {metadata['short_title']}

**Manuscript type:** {metadata['article_type']}

**Main-text word count, Introduction through Discussion:** {word_count}

## Authors

{chr(10).join(author_lines)}

## Affiliations

{chr(10).join(affiliation_lines)}

## Originating institution

{institution['name']}, {institution['street_address']}, {institution['city']}
{institution['postal_code']}, {institution['country']}

## Corresponding author

{corresponding['name']}, {corresponding['degrees']}; {corresponding['telephone']};
{corresponding['email']}; {corresponding['street_address']},
{corresponding['city']} {corresponding['postal_code']},
{corresponding['country']}

## Funding and funder role

{metadata['funding_statement']} {metadata['funder_role']}

## Conflicts of interest and industry support

{metadata['conflict_of_interest_statement']}
{metadata['industry_support_statement']}

## Acknowledgments

{metadata['acknowledgments']}

## CRediT author contributions

{contributions}

## Data sharing statement

{metadata['data_sharing_statement']}

## Ethics statement

{metadata['ethics_statement']}

## Use of generative AI and AI-assisted technologies

{metadata['generative_ai_disclosure']}
"""


def final_cover_letter(metadata: dict[str, Any], interpretation: str) -> str:
    names = ", ".join(author_name(author) for author in metadata["authors"])
    corresponding = metadata["corresponding_author"]
    dual = (
        metadata.get("dual_first_authorship_explanation")
        if metadata.get("dual_first_authorship")
        else "Not requested."
    )
    return f"""# Cover Letter

**To:** Editor, *Radiology: Artificial Intelligence*

**Manuscript title:** {metadata['manuscript_title']}

**Article type:** {metadata['article_type']}

Dear Editor,

Please consider the enclosed manuscript for publication as Original Research
in *Radiology: Artificial Intelligence*. The study evaluates previously
published BU-Net components under patient-level, equal-seed, capacity- and
compute-aware development followed by one frozen independent external-testing
session. The work does not claim RES or WC as a new component. Its contribution
is controlled evidence about whether an apparent component gain persists after
capacity, compute, stochasticity, lesion-failure, and resource controls.

The audited primary evidence statement is: {interpretation}

## Required declarations

**Complete author list:** {names}

**Subject overlap:** {metadata['subject_overlap_statement']}

**Conflicts and industry support:** {metadata['conflict_of_interest_statement']}
{metadata['industry_support_statement']}

**Sole submission:** All authors confirm sole submission to
*Radiology: Artificial Intelligence*.

**Dual first authorship or fast-track request:** {dual}

Thank you for your consideration.

Sincerely,

{corresponding['name']}, {corresponding['degrees']}

{corresponding['street_address']}, {corresponding['city']}
{corresponding['postal_code']}, {corresponding['country']}

{corresponding['email']} | {corresponding['telephone']}
"""


def set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any) -> None:
    table.autofit = False
    count = max(1, len(table.columns))
    total = 9360
    widths = [total // count] * count
    widths[-1] += total - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        cant_split.set(qn("w:val"), "true")
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[column_index]))
            tc_w.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                for run_value in paragraph.runs:
                    run_value.font.name = "Arial"
                    run_value.font.size = Pt(7.5)
                    if row_index == 0:
                        run_value.bold = True


def paginate_claim_checklist(document: Any) -> None:
    """Split the long checklist at audited row boundaries.

    LibreOffice does not reliably repeat one Word table header on every
    continuation page. Five explicit tables preserve intact rows and repeat
    the header deterministically in both DOCX and PDF.
    """
    if len(document.tables) != 1:
        raise RuntimeError("CLAIM checklist must contain exactly one source table")
    table_xml = document.tables[0]._tbl
    rows = list(table_xml.tr_lst)
    if len(rows) != 45:
        raise RuntimeError(
            f"CLAIM checklist row contract changed: expected 45, found {len(rows)}"
        )
    header = rows[0]
    data_rows = rows[1:]
    groups = (
        data_rows[0:6],
        data_rows[6:15],
        data_rows[15:25],
        data_rows[25:36],
        data_rows[36:44],
    )
    parent = table_xml.getparent()
    insertion_index = parent.index(table_xml)
    parent.remove(table_xml)
    for group_index, group in enumerate(groups):
        if group_index:
            page_break = OxmlElement("w:p")
            paragraph_properties = OxmlElement("w:pPr")
            paragraph_properties.append(OxmlElement("w:pageBreakBefore"))
            page_break.append(paragraph_properties)
            parent.insert(insertion_index, page_break)
            insertion_index += 1
        table_copy = copy.deepcopy(table_xml)
        for row in list(table_copy.tr_lst):
            table_copy.remove(row)
        header_copy = copy.deepcopy(header)
        header_properties = header_copy.find(qn("w:trPr"))
        if header_properties is not None:
            repeat = header_properties.find(qn("w:tblHeader"))
            if repeat is not None:
                header_properties.remove(repeat)
        table_copy.append(header_copy)
        for row in group:
            table_copy.append(copy.deepcopy(row))
        parent.insert(insertion_index, table_copy)
        insertion_index += 1


def style_docx(path: Path, *, preview: bool) -> None:
    document = Document(path)
    section_value = document.sections[0]
    for section_value in document.sections:
        section_value.page_width = Inches(8.5)
        section_value.page_height = Inches(11)
        section_value.orientation = WD_ORIENT.PORTRAIT
        section_value.top_margin = Inches(1)
        section_value.right_margin = Inches(1)
        section_value.bottom_margin = Inches(1)
        section_value.left_margin = Inches(1)
        section_value.header_distance = Inches(0.49)
        section_value.footer_distance = Inches(0.49)
        section_value.header.paragraphs[0].text = ""
        section_value.footer.paragraphs[0].text = ""
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    for name, size, before, after in (
        ("Title", 14, 0, 12),
        ("Heading 1", 12, 12, 6),
        ("Heading 2", 11, 10, 4),
        ("Heading 3", 11, 8, 2),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style.element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if document.paragraphs:
        first = document.paragraphs[0]
        first.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for paragraph in document.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.widow_control = True
        for run_value in paragraph.runs:
            run_value.font.name = "Arial"
            run_value._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
            run_value._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
        if preview and PREVIEW_WARNING in paragraph.text:
            paragraph.paragraph_format.line_spacing = 1.0
            for run_value in paragraph.runs:
                run_value.font.color.rgb = RGBColor(156, 0, 6)
                run_value.font.bold = True
                run_value.font.size = Pt(10)
    for table in document.tables:
        set_table_geometry(table)
    if path.stem == "06_claim_2024_checklist":
        paginate_claim_checklist(document)
    properties = document.core_properties
    properties.title = TITLE
    properties.subject = "Radiology: Artificial Intelligence submission package"
    properties.author = (
        "Author confirmation required" if preview else "Confirmed authors"
    )
    properties.comments = (
        "Pre-results preview; not for submission."
        if preview
        else "Generated from audited Gate J claims and confirmed author metadata."
    )
    document.save(path)


def audit_docx_layout(path: Path) -> dict[str, Any]:
    document = Document(path)
    sections_pass = all(
        section_value.page_width == Inches(8.5)
        and section_value.page_height == Inches(11)
        and section_value.top_margin == Inches(1)
        and section_value.right_margin == Inches(1)
        and section_value.bottom_margin == Inches(1)
        and section_value.left_margin == Inches(1)
        for section_value in document.sections
    )
    normal = document.styles["Normal"]
    normal_pass = (
        normal.font.name == "Arial"
        and normal.font.size == Pt(11)
        and normal.paragraph_format.line_spacing_rule == WD_LINE_SPACING.DOUBLE
    )
    header_footer_xml = " ".join(
        section_value.header._element.xml + section_value.footer._element.xml
        for section_value in document.sections
    )
    no_page_number_fields = (
        "PAGE" not in header_footer_xml
        and "NUMPAGES" not in header_footer_xml
    )
    tables: list[dict[str, Any]] = []
    for table in document.tables:
        table_xml = table._tbl.xml
        rows_do_not_split = all(
            row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
            for row in table.rows
        )
        tables.append(
            {
                "rows": len(table.rows),
                "columns": len(table.columns),
                "fixed_layout": 'w:type="fixed"' in table_xml,
                "rows_do_not_split": rows_do_not_split,
                "explicit_header_row": bool(table.rows),
                "merged_cells_absent": (
                    "<w:gridSpan" not in table_xml and "<w:vMerge" not in table_xml
                ),
            }
        )
    tables_pass = all(
        table["fixed_layout"]
        and table["rows_do_not_split"]
        and table["explicit_header_row"]
        and table["merged_cells_absent"]
        for table in tables
    )
    report = {
        "letter_page_and_one_inch_margins": sections_pass,
        "normal_style_arial_11_double_spaced": normal_pass,
        "page_number_fields_absent": no_page_number_fields,
        "tables": tables,
        "tables_pass": tables_pass,
    }
    report["pass"] = (
        sections_pass and normal_pass and no_page_number_fields and tables_pass
    )
    return report


def convert_markdown(source: Path, destination: Path, *, preview: bool) -> None:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("pandoc is required")
    run(
        [
            pandoc,
            str(source),
            "--from=markdown+pipe_tables+tex_math_dollars",
            "--standalone",
            "--output",
            str(destination),
        ]
    )
    style_docx(destination, preview=preview)


def convert_pdf(docx_path: Path, output_dir: Path) -> Path:
    soffice = shutil.which("soffice")
    if soffice is None:
        raise RuntimeError("LibreOffice soffice is required")
    with tempfile.TemporaryDirectory(prefix="brats-submission-lo-") as profile:
        run(
            [
                soffice,
                f"-env:UserInstallation=file://{profile}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ]
        )
    output = output_dir / f"{docx_path.stem}.pdf"
    if not output.is_file():
        raise RuntimeError(f"LibreOffice did not create {output}")
    return output


def count_words(source: str) -> int:
    source = re.sub(r"\{\{.*?\}\}", "VALUE", source)
    source = re.sub(r"<!--.*?-->", "", source, flags=re.DOTALL)
    return len(WORD.findall(source))


def main_body_word_count(source: str) -> int:
    intro = re.search(r"^## Introduction\s*$", source, re.MULTILINE)
    if intro is None:
        raise RuntimeError("Cannot identify Introduction-through-Discussion range")
    conclusion = re.search(
        r"^### Conclusion\s*$",
        source[intro.end() :],
        re.MULTILINE,
    )
    if conclusion is None:
        raise RuntimeError("Cannot identify Introduction-through-Discussion range")
    conclusion_start = intro.end() + conclusion.start()
    return count_words(source[intro.end() : conclusion_start])


def abstract_word_count(source: str) -> int:
    abstract = re.search(r"^## Abstract\s*$", source, re.MULTILINE)
    intro = re.search(r"^## Introduction\s*$", source, re.MULTILINE)
    if abstract is None or intro is None:
        raise RuntimeError("Cannot identify abstract range")
    return count_words(source[abstract.end() : intro.start()])


def source_set(mode: str, metadata: dict[str, Any] | None) -> dict[str, str]:
    preview = mode == "preview"
    main_path = (
        MANUSCRIPT / "q1q2_v2_manuscript.template.md"
        if preview
        else ROOT / "artifacts/q1q2_v2/claims/rendered_manuscript.md"
    )
    supplement_path = (
        MANUSCRIPT / "q1q2_v2_supplement.template.md"
        if preview
        else ROOT / "artifacts/q1q2_v2/claims/rendered_supplement.md"
    )
    response_path = (
        MANUSCRIPT / "q1q2_v2_response_to_reviewer.template.md"
        if preview
        else ROOT / "artifacts/q1q2_v2/claims/rendered_response_to_reviewer.md"
    )
    main = radiology_ai_manuscript(
        main_path.read_text(encoding="utf-8"), preview=preview
    )
    supplement = supplement_path.read_text(encoding="utf-8")
    response = response_path.read_text(encoding="utf-8")
    checklist = (
        SUBMISSION / "q1q2_v2_claim_2024_checklist.template.md"
    ).read_text(encoding="utf-8")
    data_statement = (
        RADAI / "data_code_availability.template.md"
    ).read_text(encoding="utf-8")
    if preview:
        supplement = add_preview_warning(replace_claims_with_pending(supplement))
        response = add_preview_warning(replace_claims_with_pending(response))
        checklist = add_preview_warning(checklist)
        data_statement = add_preview_warning(data_statement)
        cover = add_preview_warning(
            (RADAI / "cover_letter.template.md").read_text(encoding="utf-8")
        )
        title_page = add_preview_warning(
            (RADAI / "full_title_page.template.md").read_text(encoding="utf-8")
        )
    else:
        assert metadata is not None
        if UNRESOLVED.search(main + supplement + response):
            raise PermissionError("Final Gate J sources contain unresolved tokens")
        interpretation_match = re.search(
            r"The primary result was:\s*(.+?)(?:\n\n|$)",
            main,
            flags=re.DOTALL,
        )
        if interpretation_match is None:
            raise RuntimeError("Audited primary interpretation is missing")
        interpretation = " ".join(interpretation_match.group(1).split())
        title_page = final_title_page(metadata, main_body_word_count(main))
        cover = final_cover_letter(metadata, interpretation)
        data_statement = (
            "# Data and Code Availability Statement\n\n"
            + str(metadata["data_sharing_statement"]).strip()
            + "\n"
        )
        if "[FINAL PAGE/LINES" in checklist:
            raise PermissionError(
                "CLAIM checklist still lacks final page/line references"
            )
    return {
        "01_anonymized_manuscript": main,
        "02_full_title_page": title_page,
        "03_cover_letter": cover,
        "04_supplement": supplement,
        "05_response_to_reviewer": response,
        "06_claim_2024_checklist": checklist,
        "07_data_and_code_availability": data_statement,
    }


def audit_sources(sources: dict[str, str], *, mode: str) -> dict[str, Any]:
    main = sources["01_anonymized_manuscript"]
    references = section(main, "References", "Tables")
    report: dict[str, Any] = {
        "mode": mode,
        "main_text_words_introduction_through_discussion": main_body_word_count(main),
        "abstract_words": abstract_word_count(main),
        "reference_count": len(
            re.findall(r"^\d+\.\s", references, re.MULTILINE)
        ),
        "claim_tokens": sum(
            len(CLAIM_TOKEN.findall(value)) for value in sources.values()
        ),
        "unresolved_tokens": sum(
            len(UNRESOLVED.findall(value)) for value in sources.values()
        ),
        "pending_markers": sum(
            len(PENDING_MARKER.findall(value)) for value in sources.values()
        ),
    }
    report["format_limits_pass"] = (
        report["main_text_words_introduction_through_discussion"] <= 3000
        and report["abstract_words"] <= 250
        and report["reference_count"] <= 35
    )
    report["submission_ready"] = (
        mode == "final"
        and report["format_limits_pass"]
        and report["unresolved_tokens"] == 0
        and report["pending_markers"] == 0
    )
    if not report["format_limits_pass"]:
        raise RuntimeError(f"Radiology: AI word/reference limits fail: {report}")
    if mode == "final" and not report["submission_ready"]:
        raise PermissionError(f"Final source audit failed: {report}")
    return report


def build(mode: str, output_dir: Path, metadata_path: Path | None) -> dict[str, Any]:
    if mode == "final":
        require_final_gates()
        if metadata_path is None:
            raise PermissionError("--metadata is required in final mode")
        metadata = confirmed_author_metadata(metadata_path)
    else:
        metadata = None
    output_dir.mkdir(parents=True, exist_ok=True)
    sources = source_set(mode, metadata)
    audit = audit_sources(sources, mode=mode)
    files: list[dict[str, Any]] = []
    layout_audits: dict[str, dict[str, Any]] = {}
    for stem, content in sources.items():
        md_path = output_dir / f"{stem}.md"
        docx_path = output_dir / f"{stem}.docx"
        md_path.write_text(content.rstrip() + "\n", encoding="utf-8")
        convert_markdown(md_path, docx_path, preview=mode == "preview")
        layout_audits[stem] = audit_docx_layout(docx_path)
        pdf_path = convert_pdf(docx_path, output_dir)
        pdf_pages = len(PdfReader(pdf_path).pages)
        for path in (md_path, docx_path, pdf_path):
            entry = {
                "path": path.relative_to(ROOT).as_posix(),
                "sha256": digest(path),
                "bytes": path.stat().st_size,
            }
            if path.suffix == ".pdf":
                entry["pages"] = pdf_pages
            files.append(entry)
    audit["docx_layout_pass"] = all(
        report["pass"] for report in layout_audits.values()
    )
    supplement_pdf = next(
        entry
        for entry in files
        if str(entry["path"]).endswith("04_supplement.pdf")
    )
    if mode == "final" and int(supplement_pdf["pages"]) > 12:
        raise RuntimeError(
            "Radiology: AI supplement exceeds the 12-page limit: "
            f"{supplement_pdf['pages']} pages"
        )
    manifest = {
        "schema_version": 1,
        "status": "preview_not_for_submission" if mode == "preview" else "final",
        "mode": mode,
        "journal": "Radiology: Artificial Intelligence",
        "article_type": "Original Research",
        "builder_base_commit": subprocess.check_output(
            ["git", "rev-parse", "HEAD"], cwd=ROOT, text=True
        ).strip(),
        "audit": audit,
        "layout_audits": layout_audits,
        "files": sorted(files, key=lambda item: item["path"]),
        "final_blockers": (
            [
                "Gate H external results",
                "Gate I reproducibility completion",
                "Gate J claim rendering",
                "confirmed author metadata",
                "final page/line references in CLAIM",
                "final live journal-policy recheck",
            ]
            if mode == "preview"
            else []
        ),
        "visual_qa_report": (
            "submission/generated/pre_results/visual_qa.json"
            if mode == "preview"
            else None
        ),
    }
    manifest_path = output_dir / "submission_manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mode", choices=("preview", "final"), default="preview")
    parser.add_argument("--metadata", type=Path)
    parser.add_argument("--output-dir", type=Path)
    arguments = parser.parse_args()
    package_state = "pre_results" if arguments.mode == "preview" else "final"
    output_dir = arguments.output_dir or SUBMISSION / "generated" / package_state
    report = build(arguments.mode, output_dir, arguments.metadata)
    print(json.dumps(report, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
