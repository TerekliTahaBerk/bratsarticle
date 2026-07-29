#!/usr/bin/env python3
"""Build and style Gate 14 DOCX, LaTeX, and PDF deliverables."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_WIDTH = PAGE_WIDTH - 2 * MARGIN
GRAY = "E8EBEF"
DARK = RGBColor(38, 45, 55)


def _run(command: list[str]) -> None:
    subprocess.run(command, cwd=ROOT, check=True)


def _set_font(style, name: str, size: float, bold: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.color.rgb = DARK
    if bold is not None:
        font.bold = bold
    style.element.rPr.rFonts.set(qn("w:ascii"), name)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _set_cell_margins(
    cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_width(table, width_twips: int = 9360) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _add_field(run, field: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_font(normal, "Times New Roman", 11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, 16, 8),
        "Heading 2": (13, 12, 6),
        "Heading 3": (12, 8, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        style = styles[name]
        _set_font(style, "Arial", size, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    _set_font(caption, "Times New Roman", 9)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Title" not in styles:
        styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    title = styles["Title"]
    _set_font(title, "Arial", 21, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True

    if "Subtitle" not in styles:
        styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle = styles["Subtitle"]
    _set_font(subtitle, "Times New Roman", 11)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)

    for name in ("List Bullet", "List Number"):
        if name in styles:
            _set_font(styles[name], "Times New Roman", 11)
            styles[name].paragraph_format.space_after = Pt(3)


def _configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)
        section.different_first_page_header_footer = True

        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].text = ""

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "CONTROLLED EVALUATION OF BU-NET COMPONENTS"
        paragraph.style = document.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(90, 96, 105)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(8)
        _add_field(run, "PAGE")


def _style_cover(document: Document) -> None:
    title_done = False
    page_break_done = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if not title_done:
            paragraph.style = document.styles["Title"]
            title_done = True
            continue
        if text.startswith("Taha Berk Terekli") or text.startswith("^"):
            paragraph.style = document.styles["Subtitle"]
        elif (
            text.startswith("Corresponding author:")
            or text.startswith("Running title:")
            or text.startswith("Article type:")
        ):
            paragraph.style = document.styles["Subtitle"]
        if "Article type:" in text and not page_break_done:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            page_break_done = True


def _style_paragraphs(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(("Figure ", "Supplementary Figure ")):
            paragraph.style = document.styles["Caption"]
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.widow_control = True


def _style_tables(document: Document) -> None:
    for table in document.tables:
        _set_table_width(table)
        if table.rows:
            _set_repeat_table_header(table.rows[0])
        keep_table_together = len(table.rows) <= 8
        column_count = len(table.columns)
        if column_count:
            widths = [CONTENT_WIDTH / column_count] * column_count
            if column_count >= 6:
                widths[0] = Inches(1.45)
                remaining = CONTENT_WIDTH - widths[0]
                widths[1:] = [remaining / (column_count - 1)] * (column_count - 1)
            for row_index, row in enumerate(table.rows):
                _cant_split(row)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row.height = Inches(0.26)
                for column_index, cell in enumerate(row.cells):
                    cell.width = widths[column_index]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    _set_cell_margins(cell)
                    if row_index == 0:
                        _shade_cell(cell, GRAY)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        if keep_table_together and row_index < len(table.rows) - 1:
                            paragraph.paragraph_format.keep_with_next = True
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(8)
                            if row_index == 0:
                                run.font.bold = True


def _resize_images(document: Document) -> None:
    max_width = Inches(6.2)
    max_height = Inches(7.9)
    for shape in document.inline_shapes:
        ratio = min(
            1.0,
            max_width / shape.width if shape.width else 1.0,
            max_height / shape.height if shape.height else 1.0,
        )
        shape.width = int(shape.width * ratio)
        shape.height = int(shape.height * ratio)
        parent = shape._inline.getparent()
        while parent is not None and parent.tag != qn("w:p"):
            parent = parent.getparent()
        if parent is not None:
            paragraph = next(
                (item for item in document.paragraphs if item._p is parent),
                None,
            )
            if paragraph is not None:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.keep_with_next = True
                paragraph.paragraph_format.space_before = Pt(6)


def _set_core_properties(document: Document, title: str) -> None:
    properties = document.core_properties
    properties.title = title
    properties.subject = "Original research - methodological evaluation"
    properties.author = (
        "Taha Berk Terekli; Livanur Mengeş; Volkan Yusuf Hal; Ali Emre Döşer"
    )
    properties.keywords = (
        "BraTS; brain tumor segmentation; U-Net; BU-Net; reproducibility"
    )
    properties.comments = (
        "Scientific values generated from tracked repository artifacts."
    )


def postprocess_docx(path: Path, title: str, manuscript: bool) -> None:
    document = Document(path)
    _configure_styles(document)
    _configure_sections(document)
    if manuscript:
        _style_cover(document)
    _style_paragraphs(document)
    _style_tables(document)
    _resize_images(document)
    _set_core_properties(document, title)
    document.save(path)


def _pandoc_base(source: Path) -> list[str]:
    pandoc = shutil.which("pandoc")
    if pandoc is None:
        raise RuntimeError("pandoc is required")
    return [
        pandoc,
        str(source),
        "--from=markdown+pipe_tables+link_attributes+tex_math_dollars+raw_tex",
        "--resource-path",
        str(ROOT),
        "--standalone",
    ]


def build_manuscript(*, compile_pdf: bool = True) -> None:
    source = MANUSCRIPT_DIR / "final_manuscript.md"
    docx = MANUSCRIPT_DIR / "final_manuscript.docx"
    tex = MANUSCRIPT_DIR / "final_manuscript.tex"
    pdf = MANUSCRIPT_DIR / "final_manuscript.pdf"
    title = (
        "Leakage-Safe Multi-Seed Evaluation of Published BU-Net Components "
        "for Resource-Constrained 2D Glioma Segmentation"
    )

    _run([*_pandoc_base(source), "--output", str(docx)])
    postprocess_docx(docx, title, manuscript=True)

    _run(
        [
            *_pandoc_base(source),
            "--output",
            str(tex),
            "--top-level-division=section",
            "-V",
            "documentclass=article",
            "-V",
            "papersize=letter",
            "-V",
            "geometry:margin=1in",
            "-V",
            "fontsize=10pt",
            "-V",
            "linestretch=1.15",
            "-V",
            "colorlinks=true",
            "-V",
            "linkcolor=black",
            "-V",
            "urlcolor=blue",
        ]
    )
    tex_content = tex.read_text(encoding="utf-8").replace(
        "{figures/final/", "{../figures/final/"
    )
    tex_content = tex_content.replace(
        "\\begin{document}",
        "\\usepackage{caption}\n\\begin{document}",
        1,
    )
    tex_content = re.sub(
        r"\\caption\{Figure [0-9]+\. ",
        r"\\caption{",
        tex_content,
    )
    tex_content = tex_content.replace(
        "\\caption{Supplementary Figure S1.",
        "\\caption*{Supplementary Figure S1.",
    )
    tex.write_text(tex_content, encoding="utf-8")
    if not compile_pdf:
        return
    tectonic = shutil.which("tectonic")
    if tectonic is None:
        raise RuntimeError("tectonic is required")
    _run(
        [
            tectonic,
            "--only-cached",
            "--keep-logs",
            "--keep-intermediates",
            "--outdir",
            str(MANUSCRIPT_DIR),
            str(tex),
        ]
    )
    produced = MANUSCRIPT_DIR / f"{tex.stem}.pdf"
    if produced != pdf:
        produced.replace(pdf)


def build_response(*, compile_pdf: bool = True) -> None:
    source = MANUSCRIPT_DIR / "response_to_reviewer.md"
    docx = MANUSCRIPT_DIR / "response_to_reviewer.docx"
    tex = MANUSCRIPT_DIR / "response_to_reviewer.tex"
    pdf = MANUSCRIPT_DIR / "response_to_reviewer.pdf"
    title = "Response to the reviewer"
    _run([*_pandoc_base(source), "--output", str(docx)])
    postprocess_docx(docx, title, manuscript=False)
    _run(
        [
            *_pandoc_base(source),
            "--output",
            str(tex),
            "--top-level-division=section",
            "-V",
            "papersize=letter",
            "-V",
            "geometry:margin=1in",
            "-V",
            "fontsize=10pt",
            "-V",
            "colorlinks=true",
            "-V",
            "urlcolor=blue",
        ]
    )
    if not compile_pdf:
        return
    tectonic = shutil.which("tectonic")
    if tectonic is None:
        raise RuntimeError("tectonic is required")
    _run(
        [
            tectonic,
            "--only-cached",
            "--keep-logs",
            "--outdir",
            str(MANUSCRIPT_DIR),
            str(tex),
        ]
    )
    produced = MANUSCRIPT_DIR / f"{tex.stem}.pdf"
    if produced != pdf:
        produced.replace(pdf)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--target",
        choices=("all", "manuscript", "response"),
        default="all",
    )
    parser.add_argument(
        "--skip-pdf",
        action="store_true",
        help="Build DOCX and TeX only; compile PDFs separately with Tectonic.",
    )
    args = parser.parse_args()
    if args.target in ("all", "manuscript"):
        build_manuscript(compile_pdf=not args.skip_pdf)
    if args.target in ("all", "response"):
        build_response(compile_pdf=not args.skip_pdf)


if __name__ == "__main__":
    main()
